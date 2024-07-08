from openai import OpenAI
import requests
from docx import Document, Inches
import os

# Initialize the OpenAI client with your API key
client = OpenAI(
    api_key=os.environ["OPENAI_API_KEY"]
)

def generate_image(prompt):
    response = client.images.generate(
        model="dall-e-3",
        prompt=prompt,
        size="1024x1024",
        n=1,
    )
    return response.data[0].url

def save_image_to_doc(image_url, doc_path):
    # Create a new Document
    doc = Document()

    # Add a title
    doc.add_heading('Generated Image', level=1)

    # Download the image
    image_response = requests.get(image_url)
    image_path = 'generated_image.png'
    with open(image_path, 'wb') as f:
        f.write(image_response.content)

    # Add the image to the document
    doc.add_picture(image_path, width=Inches(5.0))

    # Save the document
    doc.save(doc_path)

if __name__ == "__main__":
    while True:
        user_input = input("Enter a prompt for image generation (or 'quit' to exit): ")
        if user_input.lower() in ["quit", "break"]:
            break

        # Generate the image
        image_url = generate_image(user_input)
        print("Image URL: ", image_url)

        # Save the image to a new Word document
        doc_path = 'generated_image.docx'
        save_image_to_doc(image_url, doc_path)
        print(f"Image saved to {doc_path}")
