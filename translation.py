from openai import OpenAI
from pydub import AudioSegment
import math
from docx import Document

# Initialize the client with your API key
client = OpenAI(
    api_key='sk-proj-NxQYACsIlM8gR8sha6abT3BlbkFJuHDZjuAlmedrHjJRg6Wa'
)

def split_audio(file_path, chunk_length_sec=60):  # Adjust chunk length as needed
    audio = AudioSegment.from_file(file_path)
    chunks = []
    for i in range(0, math.ceil(len(audio) / 1000 / chunk_length_sec)):
        chunks.append(audio[i * chunk_length_sec * 1000 : (i + 1) * chunk_length_sec * 1000])
    return chunks

def transcribe_audio_chunk(audio_chunk, chunk_index):
    temp_file_name = f"temp_chunk_{chunk_index}.wav"
    audio_chunk.export(temp_file_name, format="wav")
    with open(temp_file_name, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file
        )
    return transcription.text

def transcribe_audio(audio_file_path):
    chunks = split_audio(audio_file_path)
    transcriptions = [transcribe_audio_chunk(chunk, i) for i, chunk in enumerate(chunks)]
    return " ".join(transcriptions)

def translate_audio_chunk(audio_chunk, chunk_index):
    temp_file_name = f"temp_chunk_{chunk_index}.wav"
    audio_chunk.export(temp_file_name, format="wav")
    with open(temp_file_name, 'rb') as audio_file:
        translation = client.audio.translations.create(
            model="whisper-1",
            file=audio_file
        )
    return translation.text

def translate_audio(audio_file_path):
    chunks = split_audio(audio_file_path)
    translations = [translate_audio_chunk(chunk, i) for i, chunk in enumerate(chunks)]
    return " ".join(translations)

def save_as_docx(translated_text, filename):
    doc = Document()
    doc.add_heading("Translation", level=1)
    doc.add_paragraph(translated_text)
    doc.save(filename)

audio_file_path = "input-data/Balleilakka.mp4"

# First transcribe the audio
transcription = transcribe_audio(audio_file_path)
print("Transcription:", transcription)

# Then translate the audio directly
translated_text = translate_audio(audio_file_path)
print("Translation:", translated_text)

save_as_docx(translated_text, 'translation.docx')
